"""
ingest.py — Document parsing pipeline

Handles: PDF, DOCX, XLSX, Images
Steps:
  1. Detect file type
  2. Extract text (+ OCR for images/scanned PDFs)
  3. Extract tables (PDFs via camelot)
  4. Describe images (via Ollama Llama 3 vision)
  5. Chunk text
  6. Embed and store in ChromaDB
"""

import os
import base64
import requests
from pathlib import Path

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_chroma import Chroma
from langchain_core.documents import Document

CHROMA_DIR = "./vectorstore"
CHUNK_SIZE = 1000
CHUNK_OVERLAP = 200
EMBED_MODEL = "sentence-transformers/all-MiniLM-L6-v2"  # local, no API key needed
OLLAMA_URL = "http://localhost:11434"
OLLAMA_MODEL = "llama3"

# Cached embeddings instance
_embeddings = None


def _get_embeddings() -> HuggingFaceEmbeddings:
    global _embeddings
    if _embeddings is None:
        _embeddings = HuggingFaceEmbeddings(model_name=EMBED_MODEL)
    return _embeddings


def get_vectorstore():
    return Chroma(
        persist_directory=CHROMA_DIR,
        embedding_function=_get_embeddings(),
    )


def ingest_document(file_path: str, content_type: str) -> dict:
    """Main entry point. Returns stats about what was ingested."""

    docs = []
    tables_count = 0
    images_count = 0

    ext = Path(file_path).suffix.lower()

    if ext == ".pdf":
        docs, tables_count, images_count = _ingest_pdf(file_path)
    elif ext == ".docx":
        docs = _ingest_docx(file_path)
    elif ext in (".xlsx", ".xls"):
        docs = _ingest_xlsx(file_path)
    elif ext in (".png", ".jpg", ".jpeg"):
        text = _ocr_image(file_path)
        description = _describe_image(file_path)
        combined = f"[OCR Text]\n{text}\n\n[Visual Description]\n{description}"
        docs = [Document(page_content=combined, metadata={"source": file_path, "type": "image"})]
        images_count = 1
    else:
        raise ValueError(f"Unsupported file type: {ext}")

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
    )
    chunks = splitter.split_documents(docs)

    filename = Path(file_path).name
    for chunk in chunks:
        chunk.metadata.setdefault("filename", filename)

    db = get_vectorstore()
    db.add_documents(chunks)

    return {
        "chunks": len(chunks),
        "tables": tables_count,
        "images": images_count,
    }


# ── PDF ──────────────────────────────────────────────

def _ingest_pdf(file_path: str):
    from langchain_community.document_loaders import PyPDFLoader

    loader = PyPDFLoader(file_path)
    pages = loader.load()

    docs = []
    for page in pages:
        if len(page.page_content.strip()) < 50:
            ocr_text = _ocr_pdf_page(file_path, page.metadata.get("page", 0))
            page.page_content = ocr_text
        docs.append(page)

    tables_count = 0
    try:
        import camelot
        tables = camelot.read_pdf(file_path, pages="all")
        for i, table in enumerate(tables):
            df = table.df
            table_text = df.to_markdown(index=False)
            table_doc = Document(
                page_content=f"[Table {i+1}]\n{table_text}",
                metadata={"source": file_path, "type": "table", "table_index": i},
            )
            docs.append(table_doc)
            tables_count += 1
    except Exception:
        pass

    images_count = 0
    try:
        import fitz
        pdf = fitz.open(file_path)
        for page_num, page in enumerate(pdf):
            for img_index, img in enumerate(page.get_images()):
                xref = img[0]
                base_image = pdf.extract_image(xref)
                image_bytes = base_image["image"]
                desc = _describe_image_bytes(image_bytes)
                img_doc = Document(
                    page_content=f"[Image on page {page_num+1}]\n{desc}",
                    metadata={"source": file_path, "type": "image", "page": page_num + 1},
                )
                docs.append(img_doc)
                images_count += 1
                if images_count >= 5:
                    break
    except Exception:
        pass

    return docs, tables_count, images_count


def _ocr_pdf_page(file_path: str, page_number: int) -> str:
    try:
        import fitz
        from PIL import Image
        import pytesseract
        import io

        pdf = fitz.open(file_path)
        page = pdf[page_number]
        pix = page.get_pixmap(dpi=200)
        img = Image.open(io.BytesIO(pix.tobytes("png")))
        return pytesseract.image_to_string(img)
    except Exception as e:
        return f"[OCR failed: {e}]"


# ── DOCX ──────────────────────────────────────────────

def _ingest_docx(file_path: str) -> list:
    from docx import Document as DocxDocument

    doc = DocxDocument(file_path)
    full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

    table_texts = []
    for i, table in enumerate(doc.tables):
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        table_md = "\n".join([" | ".join(row) for row in rows])
        table_texts.append(f"[Table {i+1}]\n{table_md}")

    combined = full_text
    if table_texts:
        combined += "\n\n" + "\n\n".join(table_texts)

    return [Document(page_content=combined, metadata={"source": file_path, "type": "docx"})]


# ── XLSX ──────────────────────────────────────────────

def _ingest_xlsx(file_path: str) -> list:
    import openpyxl

    wb = openpyxl.load_workbook(file_path)
    docs = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue

        headers = [str(c) if c is not None else "" for c in rows[0]]
        lines = ["| " + " | ".join(headers) + " |"]
        lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
        for row in rows[1:]:
            cells = [str(c) if c is not None else "" for c in row]
            lines.append("| " + " | ".join(cells) + " |")

        sheet_text = f"[Sheet: {sheet_name}]\n" + "\n".join(lines)
        docs.append(Document(
            page_content=sheet_text,
            metadata={"source": file_path, "type": "xlsx", "sheet": sheet_name},
        ))

    return docs


# ── Image helpers ──────────────────────────────────────

def _ocr_image(file_path: str) -> str:
    try:
        import pytesseract
        from PIL import Image
        img = Image.open(file_path)
        return pytesseract.image_to_string(img)
    except Exception as e:
        return f"[OCR failed: {e}]"


def _describe_image(file_path: str) -> str:
    with open(file_path, "rb") as f:
        return _describe_image_bytes(f.read())


def _describe_image_bytes(image_bytes: bytes) -> str:
    """Use Ollama Llama 3 vision to describe an image."""
    try:
        b64 = base64.b64encode(image_bytes).decode("utf-8")
        payload = {
            "model": OLLAMA_MODEL,
            "messages": [
                {
                    "role": "user",
                    "content": (
                        "Describe this image in detail. "
                        "Focus on any text, data, charts, diagrams, or key visual information."
                    ),
                    "images": [b64],
                }
            ],
            "stream": False,
        }
        resp = requests.post(f"{OLLAMA_URL}/api/chat", json=payload, timeout=120)
        resp.raise_for_status()
        return resp.json()["message"]["content"]
    except Exception as e:
        return f"[Image description unavailable: {e}]"
